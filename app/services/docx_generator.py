from __future__ import annotations

import io
from pathlib import Path

from docx import Document

from app.models.schemas import ALL_DAYS, CleaningPlan


def plan_to_docx_bytes(plan: CleaningPlan) -> bytes:
    document = Document()
    heading = "Renholdsplan"
    if plan.template_name:
        heading = f"{heading} – {plan.template_name}"
    document.add_heading(heading, level=1)
    document.add_paragraph(f"Totalt dekket areal: {plan.total_area_m2:.0f} m²")

    table = document.add_table(rows=1, cols=10)
    headers = ["AREAL", "BESKRIVELSE", "ETG"] + ALL_DAYS
    hdr_cells = table.rows[0].cells
    for idx, label in enumerate(headers):
        hdr_cells[idx].text = label

    for entry in plan.entries:
        row = table.add_row().cells
        row[0].text = f"{entry.room_name} ({entry.area_m2 or '-'} m²)"
        row[1].text = entry.description
        row[2].text = entry.floor or "-"
        for day_idx, day in enumerate(ALL_DAYS, start=3):
            row[day_idx].text = "X" if entry.frequency.get(day, False) else ""

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
